import os
import subprocess
import sys
import shutil

def install_deps():
    subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", "python-docx", "reportlab"])

def generate_docx():
    from docx import Document
    doc = Document()
    doc.add_heading('Srpska Priča (Test DOCX)', 0)
    
    p1 = doc.add_paragraph()
    p1.add_run('Mali dečak i lepa knjiga. Dečak čita novu knjigu svaki dan.')
    
    p2 = doc.add_paragraph()
    p2.add_run('On voli da uči srpski jezik u školi. Njegov prijatelj takođe uči jezik i radi u gradu.')
    
    p3 = doc.add_paragraph()
    p3.add_run('Oni govore srpski veoma dobro. Danas je lep dan. Moći ćemo da radimo i čitamo u kući.')
    
    doc.save('test_story.docx')
    print("test_story.docx generated.")
    
    dest = os.path.join('frontend', 'assets', 'test_story.docx')
    shutil.copy2('test_story.docx', dest)
    print(f"Copied test_story.docx to {dest}")

def generate_pdf():
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    font_path = os.path.join('frontend', 'assets', 'fonts', 'NotoSans-Regular.ttf')
    font_bold_path = os.path.join('frontend', 'assets', 'fonts', 'NotoSans-Bold.ttf')
    
    pdfmetrics.registerFont(TTFont('NotoSans', font_path))
    pdfmetrics.registerFont(TTFont('NotoSans-Bold', font_bold_path))
    
    doc = SimpleDocTemplate("test_story.pdf", pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Title'],
        fontName='NotoSans-Bold',
        fontSize=24,
        leading=28
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontName='NotoSans',
        fontSize=12,
        leading=16
    )
    
    story.append(Paragraph("Srpska Priča (Test PDF)", title_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph("Mali dečak i lepa knjiga. Dečak čita novu knjigu каждый день. Dečak čita novu knjigu svaki dan.", body_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("On voli da uči srpski jezik u školi. Njegov prijatelj takođe uči jezik i radi u gradu.", body_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Oni govore srpski veoma dobro. Danas je lep dan. Moći ćemo da radimo i čitamo u kući.", body_style))
    
    doc.build(story)
    print("test_story.pdf generated.")
    
    dest = os.path.join('frontend', 'assets', 'test_story.pdf')
    shutil.copy2('test_story.pdf', dest)
    print(f"Copied test_story.pdf to {dest}")

if __name__ == "__main__":
    try:
        install_deps()
        generate_docx()
        generate_pdf()
    except Exception as e:
        print(f"Error generating test files: {e}")
