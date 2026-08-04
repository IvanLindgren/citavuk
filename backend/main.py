import asyncio
import io
import os
import sqlite3
import hashlib
from pathlib import Path
from contextlib import asynccontextmanager
from fastapi import Depends, FastAPI, File, HTTPException, Query, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import Dict, Optional
import logging

try:
    from .redis_cache import redis_cache
    from .urlguard import UrlRejected, check_url, open_checked
    from .ratelimit import make_limiter
    from . import analysis_cache
except ImportError:
    from redis_cache import redis_cache
    from urlguard import UrlRejected, check_url, open_checked
    from ratelimit import make_limiter
    import analysis_cache

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

import threading
import time
import urllib.request

# Global NLP pipeline (loaded once on startup via lifespan)
nlp = None
nlp_active = False

def ping_self():
    space_host = os.environ.get("SPACE_HOST")
    if not space_host:
        logging.info("No SPACE_HOST env var found, self-ping disabled.")
        return
    
    url = f"https://{space_host}/health"
    logging.info(f"Started self-ping thread for {url} every 12 hours.")
    while True:
        # Sleep for 12 hours
        time.sleep(12 * 3600)
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "Chitavuk-KeepAlive/1.0"})
            with urllib.request.urlopen(req, timeout=10):
                logging.info(f"Self-ping successful: {url}")
        except Exception as e:
            logging.error(f"Self-ping failed: {e}")


@asynccontextmanager
async def lifespan(app: FastAPI):
    global nlp, nlp_active
    threading.Thread(target=ping_self, daemon=True).start()
    if redis_cache.configured:
        if redis_cache.ping():
            logging.info("Redis cache connected.")
        else:
            logging.warning("Redis configured but unavailable; local caches remain active.")
    try:
        import classla
        logging.info("Downloading/verifying Serbian models...")
        classla.download('sr')
        logging.info("Initializing CLASSLA Serbian pipeline...")
        nlp = classla.Pipeline('sr', processors='tokenize,pos,lemma')
        nlp_active = True
        logging.info("CLASSLA initialized.")
    except Exception as e:
        logging.warning(f"CLASSLA unavailable ({e}); using SQLite-only analysis.")
        nlp_active = False
    yield


# Swagger в проде закрыт — схема API посторонним не нужна.
# Локально включить: CITAVUK_ENABLE_DOCS=1.
if os.environ.get("CITAVUK_ENABLE_DOCS") == "1":
    app = FastAPI(title="Serbian NLP Backend", lifespan=lifespan)
else:
    app = FastAPI(
        title="Serbian NLP Backend",
        lifespan=lifespan,
        docs_url=None,
        redoc_url=None,
        openapi_url=None,
    )
# Разрешаем доступ из Flutter (web — другой origin/порт; desktop/mobile — не мешает).
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Per-IP лимиты тяжёлых эндпоинтов (rate/мин, стартовый burst).
_analyze_limit = make_limiter(30, 10)
_extract_limit = make_limiter(10, 3)
_translate_limit = make_limiter(30, 10)
_img_limit = make_limiter(60, 20)
_article_limit = make_limiter(20, 5)
_tts_limit = make_limiter(30, 10)

# ---------------------------------------------------------------------------
# Извлечение документов. Обычные PDF и DOCX веб-клиент разбирает локально,
# сюда попадают прежде всего PDF-сканы без текстового слоя.
# ---------------------------------------------------------------------------

_DOCUMENT_MAX_BYTES = 32 * 1024 * 1024
_OCR_MAX_PAGES = 80
_TESSDATA_DIR = Path(os.environ.get("CITAVUK_TESSDATA_DIR", "/tmp/citavuk-tessdata"))
_TESSDATA_MODELS = ("srp", "srp_latn")
_TESSDATA_BASE = (
    "https://raw.githubusercontent.com/tesseract-ocr/tessdata_fast/main"
)


def _ensure_serbian_tessdata() -> Path:
    """Загружает компактные официальные модели Tesseract при первом OCR."""
    _TESSDATA_DIR.mkdir(parents=True, exist_ok=True)
    for model in _TESSDATA_MODELS:
        target = _TESSDATA_DIR / f"{model}.traineddata"
        if target.exists() and target.stat().st_size > 100_000:
            continue
        temporary = target.with_suffix(".tmp")
        urllib.request.urlretrieve(
            f"{_TESSDATA_BASE}/{model}.traineddata",
            temporary,
        )
        temporary.replace(target)
    return _TESSDATA_DIR


def _document_letter_count(text: str) -> int:
    return sum(character.isalpha() for character in text)


def _clean_document_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _extract_docx_bytes(data: bytes) -> dict:
    from docx import Document

    document = Document(io.BytesIO(data))
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    text = _clean_document_text("\n\n".join(block for block in blocks if block))
    if not text:
        raise ValueError("В DOCX не нашлось текста.")
    return {
        "text": text,
        "pages": 0,
        "scannedPages": 0,
        "ocrUsed": False,
    }


def _ocr_pdf_page(page) -> str:
    from PIL import ImageOps
    import pytesseract

    tessdata = _ensure_serbian_tessdata()
    bitmap = page.render(scale=2.4)
    try:
        image = ImageOps.autocontrast(bitmap.to_pil().convert("L"))
        return pytesseract.image_to_string(
            image,
            lang="srp+srp_latn",
            config=f'--oem 1 --psm 6 --tessdata-dir "{tessdata}"',
        )
    finally:
        bitmap.close()


def _extract_pdf_bytes(data: bytes) -> dict:
    import pypdfium2 as pdfium

    document = pdfium.PdfDocument(data)
    native_pages: list[str] = []
    scanned_indexes: list[int] = []
    try:
        for index in range(len(document)):
            page = document[index]
            text_page = page.get_textpage()
            try:
                text = _clean_document_text(text_page.get_text_range())
            finally:
                text_page.close()
                page.close()
            native_pages.append(text)
            if _document_letter_count(text) < 30:
                scanned_indexes.append(index)

        if len(scanned_indexes) > _OCR_MAX_PAGES:
            raise ValueError(
                "В PDF слишком много страниц-сканов для бесплатного OCR "
                f"({len(scanned_indexes)}, максимум {_OCR_MAX_PAGES}). "
                "Разделите файл на части."
            )

        for index in scanned_indexes:
            page = document[index]
            try:
                native_pages[index] = _clean_document_text(_ocr_pdf_page(page))
            finally:
                page.close()
    finally:
        document.close()

    text = _clean_document_text(
        "\n\n".join(page for page in native_pages if page.strip())
    )
    if not text:
        raise ValueError("Не удалось распознать текст в PDF.")
    return {
        "text": text,
        "pages": len(native_pages),
        "scannedPages": len(scanned_indexes),
        "ocrUsed": bool(scanned_indexes),
    }


def _extract_document_bytes(filename: str, data: bytes) -> dict:
    lowered = filename.lower()
    if data.startswith(b"%PDF-") or lowered.endswith(".pdf"):
        return _extract_pdf_bytes(data)
    if data.startswith(b"PK") or lowered.endswith(".docx"):
        return _extract_docx_bytes(data)
    raise ValueError("Поддерживаются только PDF и DOCX.")


@app.post("/documents/extract")
async def extract_document(
    file: UploadFile = File(...),
    _: None = Depends(_extract_limit),
):
    data = await file.read(_DOCUMENT_MAX_BYTES + 1)
    if len(data) > _DOCUMENT_MAX_BYTES:
        raise HTTPException(
            status_code=413,
            detail="Файл больше 32 МБ. Разделите его на несколько частей.",
        )
    if not data:
        raise HTTPException(status_code=400, detail="Файл пуст.")

    cache_key = hashlib.sha256(data).hexdigest()
    cached = redis_cache.get_json(f"document:{cache_key}")
    if isinstance(cached, dict) and cached.get("text"):
        return cached

    try:
        result = await asyncio.to_thread(
            _extract_document_bytes,
            file.filename or "document",
            data,
        )
    except ValueError as error:
        raise HTTPException(status_code=422, detail=str(error)) from error
    except Exception as error:
        logging.exception("Document extraction failed")
        raise HTTPException(
            status_code=500,
            detail="Не удалось извлечь текст из документа.",
        ) from error

    redis_cache.set_json(f"document:{cache_key}", result, 7 * 24 * 3600)
    return result


@app.get("/health")
def health():
    return {
        "status": "ok",
        "nlp": nlp_active,
        "redis": redis_cache.ping() if redis_cache.configured else False,
    }

# DB path check
DB_PATH = "lexicon.db"
if not os.path.exists(DB_PATH):
    # Fallback to parent dir if run from backend/
    DB_PATH = "../lexicon.db"

# Global read-only connection
_db_conn = None

def get_db():
    global _db_conn
    if _db_conn is None and os.path.exists(DB_PATH):
        _db_conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    return _db_conn

# Transliteration mappings
CYR_TO_LAT = {
    'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'ђ': 'đ', 'е': 'e', 'ж': 'ž',
    'з': 'z', 'и': 'i', 'ј': 'j', 'к': 'k', 'л': 'l', 'љ': 'lj', 'м': 'm', 'н': 'n',
    'њ': 'nj', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'ћ': 'ć', 'у': 'u',
    'ф': 'f', 'х': 'h', 'ц': 'c', 'ч': 'č', 'џ': 'dž', 'ш': 'š',
    'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'G', 'Д': 'D', 'Ђ': 'Đ', 'Е': 'E', 'Ж': 'Ž',
    'З': 'Z', 'И': 'I', 'Ј': 'J', 'К': 'K', 'Л': 'L', 'Љ': 'Lj', 'М': 'M', 'Н': 'N',
    'Њ': 'Nj', 'О': 'O', 'П': 'P', 'Р': 'R', 'С': 'S', 'Т': 'T', 'Ћ': 'Ć', 'У': 'U',
    'Ф': 'F', 'Х': 'H', 'Ц': 'C', 'Ч': 'Č', 'Џ': 'Dž', 'Ш': 'Š'
}

def to_latin(text: str) -> str:
    # Transliterates Serbian Cyrillic to Latin
    res = []
    i = 0
    while i < len(text):
        char = text[i]
        res.append(CYR_TO_LAT.get(char, char))
        i += 1
    return "".join(res)

class AnalyzeRequest(BaseModel):
    sentence: str = Field(max_length=2000)
    start_offset: int
    end_offset: int
    token_text: str = Field(max_length=100)

def _parse_feats(s: str) -> Dict[str, str]:
    d: Dict[str, str] = {}
    if s and s != "_":
        for part in s.split("|"):
            if "=" in part:
                k, v = part.split("=", 1)
                d[k] = v
    return d


def query_local_lexicon(word_latin: str) -> list:
    """Returns lexicon rows (lemma, upos, feats, msd) for a word form."""
    conn = get_db()
    if not conn:
        return []
    cursor = conn.cursor()
    cursor.execute(
        "SELECT lemma, upos, feats, msd FROM lexicon WHERE form = ?",
        (word_latin.lower(),),
    )
    return cursor.fetchall()

def get_forms_from_lexicon(lemma: str, upos: str) -> Dict[str, str]:
    """Retrieves standard forms for a lemma using UD feats from local DB."""
    forms: Dict[str, str] = {}
    conn = get_db()
    if not conn:
        return forms
    cursor = conn.cursor()
    cursor.execute("SELECT form, feats FROM lexicon WHERE lemma = ?", (lemma.lower(),))
    rows = cursor.fetchall()

    for form, feats_str in rows:
        f = _parse_feats(feats_str)
        if upos in ("VERB", "AUX"):
            if f.get("VerbForm") == "Inf":
                forms["infinitive"] = form
            elif f.get("Tense") == "Pres" and f.get("Person") == "1" and f.get("Number") == "Sing":
                forms["present_1sg"] = form
        elif upos in ("NOUN", "PROPN"):
            if f.get("Case") == "Nom" and f.get("Number") == "Sing":
                forms["nominative_singular"] = form
            elif f.get("Case") == "Nom" and f.get("Number") == "Plur":
                forms["nominative_plural"] = form
        elif upos == "ADJ":
            if f.get("Case") == "Nom" and f.get("Number") == "Sing":
                g = f.get("Gender")
                if g == "Masc":
                    forms["nominative_masculine"] = form
                elif g == "Fem":
                    forms["nominative_feminine"] = form
                elif g == "Neut":
                    forms["nominative_neuter"] = form

    return forms

def get_dictionary_translation(word: str, lemma: str) -> Optional[str]:
    """Looks up translation in SQLite database."""
    conn = get_db()
    if not conn:
        return None
    cursor = conn.cursor()
    
    # Try exact word
    cursor.execute("SELECT translation FROM dictionary WHERE word = ?", (word.lower(),))
    row = cursor.fetchone()
    if row:
        return row[0]
        
    # Try lemma
    cursor.execute("SELECT translation FROM dictionary WHERE word = ?", (lemma.lower(),))
    row = cursor.fetchone()
    if row:
        return row[0]
    return None

def fetch_online_translation(word: str) -> str:
    """Fallback translator using Google Translate."""
    try:
        from deep_translator import GoogleTranslator
        translated = GoogleTranslator(source='sr', target='ru').translate(word)
        return translated
    except Exception as e:
        logging.error(f"Translation service failed: {e}")
        return "[Перевод недоступен]"

@app.post("/analyze")
def analyze_token(req: AnalyzeRequest, _: None = Depends(_analyze_limit)):
    analysis_key = hashlib.sha256(
        (
            f"{req.sentence}\0{req.start_offset}\0"
            f"{req.end_offset}\0{req.token_text}"
        ).encode("utf-8")
    ).hexdigest()
    cached_analysis = redis_cache.get_json(f"analysis:{analysis_key}")
    if isinstance(cached_analysis, dict):
        return cached_analysis
    # L2: устойчивый SQLite-кеш, переживает перезапуск и отсутствие Redis.
    cached_analysis = analysis_cache.get(analysis_key)
    if isinstance(cached_analysis, dict):
        redis_cache.set_json(f"analysis:{analysis_key}", cached_analysis, 24 * 3600)
        return cached_analysis

    word_lat = to_latin(req.token_text).strip()
    
    # Fast path for multi-word phrases
    if " " in word_lat:
        translation = fetch_online_translation(req.token_text)
        result = {
            "lemma": word_lat,
            "upos": "PHRASE",
            "feats": {},
            "forms": {},
            "translation": translation
        }
        redis_cache.set_json(f"analysis:{analysis_key}", result, 24 * 3600)
        analysis_cache.set(analysis_key, result)
        return result
        
    lemma = word_lat.lower()
    upos = "UNKNOWN"
    feats = {}
    
    # Engine A: CLASSLA (Contextual NLP)
    if nlp_active and nlp:
        try:
            doc = nlp(req.sentence)
            matched_word = None
            
            # Find the word that corresponds to the character offsets
            for sentence in doc.sentences:
                for word in sentence.words:
                    # Stanza stores character offset inside word.misc as 'start_char=X|end_char=Y'
                    misc = word.misc if hasattr(word, 'misc') else ""
                    start_c = None
                    end_c = None
                    if misc:
                        parts = misc.split('|')
                        for p in parts:
                            if p.startswith('start_char='):
                                start_c = int(p.split('=')[1])
                            elif p.startswith('end_char='):
                                end_c = int(p.split('=')[1])
                                
                    # If offsets overlap or text matches
                    if (start_c is not None and end_c is not None and 
                            start_c <= req.start_offset and end_c >= req.end_offset):
                        matched_word = word
                        break
                    elif to_latin(word.text).lower() == word_lat.lower():
                        matched_word = word
            
            if matched_word:
                lemma = to_latin(matched_word.lemma)
                upos = matched_word.upos
                # Parse features
                if matched_word.feats:
                    # Feats string format: "Case=Nom|Gender=Masc|Number=Sing"
                    feat_parts = matched_word.feats.split('|')
                    for part in feat_parts:
                        if '=' in part:
                            k, v = part.split('=', 1)
                            feats[k] = v
        except Exception as e:
            logging.error(f"CLASSLA analysis failed: {e}")
            # Fall back to SQLite

    # Engine B: SQLite fallback (Lexicon lookup with UD upos/feats)
    if upos == "UNKNOWN":
        rows = query_local_lexicon(word_lat)
        if rows:
            content = {"NOUN", "VERB", "ADJ", "PROPN", "ADV", "NUM"}
            best = next((r for r in rows if r[1] in content), rows[0])
            lemma = best[0]
            upos = best[1] or "UNKNOWN"
            feats = _parse_feats(best[2])

    # Engine C: Online Wiktionary fallback
    if upos == "UNKNOWN":
        online = fetch_wiktionary_lemma(word_lat)
        if online:
            lemma, upos = online

    # Get inflected forms from local lexicon DB
    forms = get_forms_from_lexicon(lemma, upos)
    
    # Translate (SQLite Dictionary -> DeepL/Google fallback)
    translation = get_dictionary_translation(req.token_text, lemma)
    if not translation:
        translation = fetch_online_translation(req.token_text)
        
    # Contextual translation logic
    contextual_translation = None
    if req.sentence and req.start_offset is not None and req.end_offset is not None:
        try:
            start = req.start_offset
            end = req.end_offset
            if 0 <= start <= len(req.sentence) and 0 <= end <= len(req.sentence) and start < end:
                tagged_sentence = req.sentence[:start] + f"<w>{req.token_text}</w>" + req.sentence[end:]
                translated_sentence = fetch_online_translation(tagged_sentence)
                import re
                match = re.search(r"<w[^>]*>(.*?)</w>", translated_sentence, re.IGNORECASE | re.DOTALL)
                if match:
                    contextual_translation = match.group(1).strip()
        except Exception as e:
            logging.error(f"Contextual translation failed: {e}")

    result = {
        "lemma": lemma,
        "upos": upos,
        "feats": feats,
        "forms": forms,
        "translation": translation,
        "contextual_translation": contextual_translation
    }
    redis_cache.set_json(f"analysis:{analysis_key}", result, 24 * 3600)
    analysis_cache.set(analysis_key, result)
    return result

import urllib.parse
import re
import ssl
import json
import html as _html

def fetch_wiktionary_lemma(word: str) -> Optional[tuple[str, str]]:
    try:
        url = "https://en.wiktionary.org/w/api.php?" + urllib.parse.urlencode({
            "action": "query",
            "titles": word,
            "prop": "revisions",
            "rvprop": "content",
            "format": "json",
            "redirects": 1
        })
        ctx = ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE
        req = urllib.request.Request(url, headers={"User-Agent": "Chitavuk/1.0 (denis@example.com)"})
        with urllib.request.urlopen(req, context=ctx, timeout=3) as response:
            data = json.loads(response.read().decode('utf-8'))
            pages = data.get("query", {}).get("pages", {})
            for page_id, page in pages.items():
                if page_id == "-1":
                    continue
                revisions = page.get("revisions", [])
                if not revisions:
                    continue
                content = revisions[0].get("*", "")
                
                if "==Serbo-Croatian==" not in content and "==Serbian==" not in content:
                    continue
                
                sc_match = re.search(r"==(?:Serbo-Croatian|Serbian)==(.*?)(?:==[a-zA-Z -]+==|$)", content, re.DOTALL)
                if not sc_match:
                    continue
                sc_content = sc_match.group(1)
                
                lemma = None
                upos = "UNKNOWN"
                
                match3 = re.search(r"\{\{(?:inflection of|infl of)\|(?:sh|sr)\|([^}|]+)", sc_content)
                match1 = re.search(r"of\s+(?:the\s+)?\[\[([^\]|]+)(?:\|[^\]]+)?\]\]", sc_content)
                match2 = re.search(r"\{\{m\|(?:sh|sr)\|([^}|]+)", sc_content)
                
                if match3:
                    lemma = match3.group(1)
                elif match1:
                    lemma = match1.group(1)
                elif match2:
                    lemma = match2.group(1)
                
                if "===Noun===" in sc_content:
                    upos = "NOUN"
                elif "===Verb===" in sc_content:
                    upos = "VERB"
                elif "===Adjective===" in sc_content:
                    upos = "ADJ"
                elif "===Adverb===" in sc_content:
                    upos = "ADV"
                elif "===Pronoun===" in sc_content:
                    upos = "PRON"
                elif "===Proper noun===" in sc_content:
                    upos = "PROPN"
                
                if lemma:
                    return lemma.strip().lower(), upos
                elif upos != "UNKNOWN":
                    return word.strip().lower(), upos
    except Exception as e:
        logging.error(f"Wiktionary lookup failed: {e}")
    return None

# ---------------------------------------------------------------------------
# Новости/статьи на сербском: RSS-ленты по темам + извлечение полного текста.
# Парсинг идёт на сервере (нет CORS, надёжные библиотеки), приложение зовёт
# /news и /article. Ленты можно править под себя.
# ---------------------------------------------------------------------------
import calendar

# Ленты проверены (отдают статьи). Темы агрегируют несколько СМИ; дубли по
# ссылке отсекаются, сортировка — по дате.
NEWS_FEEDS = {
    "general": [
        "https://n1info.rs/feed/",
        "https://www.danas.rs/feed/",
        "https://nova.rs/feed/",
    ],
    "politics": [
        "https://n1info.rs/vesti/feed/",
        "https://nova.rs/vesti/politika/feed/",
    ],
    "culture": [
        "https://www.blic.rs/rss/Kultura",
        "https://nova.rs/kultura/feed/",
        "https://n1info.rs/kultura/feed/",
    ],
    "trending": [  # «лента дня» — самое читаемое сегодня (Blic) + свежее
        "https://www.blic.rs/rss/danasnji-najcitaniji",
        "https://nova.rs/feed/",
        "https://n1info.rs/feed/",
    ],
    "science": [
        "https://naukakrozprice.rs/feed/",
        "https://nova.rs/it/feed/",
        "https://www.blic.rs/rss/IT",
    ],
}

# Кэш ленты по теме (чтобы не дёргать RSS на каждый заход).
_NEWS_CACHE = {}
_NEWS_TTL = 300  # 5 минут


def _clean_summary(html: str) -> str:
    text = re.sub(r"<[^>]+>", "", html or "")
    text = re.sub(r"\s+", " ", text).strip()
    return text[:300]


def _entry_timestamp(e) -> int:
    for key in ("published_parsed", "updated_parsed"):
        val = e.get(key)
        if val:
            try:
                return calendar.timegm(val)
            except Exception:
                pass
    return 0


def _entry_image(e) -> Optional[str]:
    # media:content / media:thumbnail
    for key in ("media_content", "media_thumbnail"):
        media = e.get(key)
        if media:
            for m in media:
                if m.get("url"):
                    return m["url"]
    # enclosure-картинки
    for link in e.get("links", []):
        if link.get("rel") == "enclosure" and str(link.get("type", "")).startswith("image"):
            return link.get("href")
    # первый <img> в summary/content
    html = e.get("summary", "") or ""
    if e.get("content"):
        try:
            html += e["content"][0].get("value", "")
        except Exception:
            pass
    m = re.search(r'<img[^>]+src=["\']([^"\']+)["\']', html)
    if m:
        return m.group(1)
    return None


@app.get("/news")
def news(topic: str = "general", limit: int = 25):
    import feedparser
    # Кэш: отдаём свежий результat, не дёргая RSS чаще раза в 5 минут.
    now = time.time()
    cached = _NEWS_CACHE.get(topic)
    if cached and now - cached[0] < _NEWS_TTL:
        return {"topic": topic, "items": cached[1][:limit], "cached": True}
    shared = redis_cache.get_json(f"news:{topic}")
    if isinstance(shared, list):
        _NEWS_CACHE[topic] = (now, shared)
        return {"topic": topic, "items": shared[:limit], "cached": True}

    feeds = NEWS_FEEDS.get(topic, NEWS_FEEDS["general"])
    items = []
    seen = set()
    for feed_url in feeds:
        try:
            d = feedparser.parse(feed_url)
            source = ""
            try:
                source = d.feed.get("title", "")
            except Exception:
                pass
            for e in d.entries:
                link = e.get("link", "")
                if not link or link in seen:
                    continue
                seen.add(link)
                items.append({
                    "title": (e.get("title", "") or "").strip(),
                    "summary": _clean_summary(e.get("summary", "")),
                    "image": _entry_image(e),
                    "source": source,
                    "link": link,
                    "published": e.get("published", "") or e.get("updated", ""),
                    "published_ts": _entry_timestamp(e),
                })
        except Exception as ex:
            logging.error(f"RSS feed failed ({feed_url}): {ex}")
    items.sort(key=lambda x: x.get("published_ts", 0), reverse=True)
    _NEWS_CACHE[topic] = (now, items)
    redis_cache.set_json(f"news:{topic}", items, _NEWS_TTL)
    return {"topic": topic, "items": items[:limit]}


_TRANSLATE_LANGS = {"sr", "ru", "en"}


@app.get("/translate")
def translate(
    q: str = Query(..., max_length=500),
    sl: str = "sr",
    tl: str = "ru",
    _: None = Depends(_translate_limit),
):
    """Перевод текста (для веба, где прямой запрос к Google блокируется CORS)."""
    if sl not in _TRANSLATE_LANGS or tl not in _TRANSLATE_LANGS:
        raise HTTPException(status_code=400, detail="Неподдерживаемый язык.")
    cache_key = hashlib.sha256(
        f"{sl.strip().lower()}:{tl.strip().lower()}:{q.strip().lower()}".encode("utf-8")
    ).hexdigest()
    cached = redis_cache.get_text(f"translation:{cache_key}")
    if cached is not None:
        return {"translation": cached, "cached": True}
    try:
        from deep_translator import GoogleTranslator
        out = GoogleTranslator(source=sl, target=tl).translate(q)
        if out:
            redis_cache.set_text(f"translation:{cache_key}", out, 24 * 3600)
        return {"translation": out or ""}
    except Exception as ex:
        logging.error(f"/translate failed: {ex}")
        return {"translation": ""}


@app.get("/img")
def img_proxy(url: str, _: None = Depends(_img_limit)):
    """Прокси картинок: грузим на сервере и отдаём приложению — чтобы новостные
    изображения работали и в вебе (нет CORS). Только хосты лент (urlguard)."""
    from fastapi.responses import Response
    try:
        check_url(url)
    except UrlRejected as ex:
        logging.warning(f"Image proxy rejected ({url}): {ex}")
        return Response(status_code=404)
    try:
        with open_checked(
            url,
            headers={"User-Agent": "Mozilla/5.0"},
            timeout=10,
        ) as r:
            content = r.read(10 * 1024 * 1024 + 1)
            ctype = r.headers.get("Content-Type", "")
        if len(content) > 10 * 1024 * 1024:
            return Response(status_code=404)
        if not ctype.lower().startswith("image/"):
            return Response(status_code=404)
        return Response(content=content, media_type=ctype)
    except Exception as ex:
        logging.error(f"Image proxy failed ({url}): {ex}")
        from fastapi.responses import Response as _R
        return _R(status_code=404)


@app.get("/article")
def article(url: str, _: None = Depends(_article_limit)):
    """Извлекает основной текст статьи и заглавную картинку по ссылке."""
    try:
        check_url(url)
    except UrlRejected as ex:
        logging.warning(f"Article rejected ({url}): {ex}")
        return {"error": "fetch_failed"}
    cache_key = hashlib.sha256(url.strip().encode("utf-8")).hexdigest()
    cached = redis_cache.get_json(f"article:{cache_key}")
    if isinstance(cached, dict) and isinstance(cached.get("paragraphs"), list):
        return cached
    try:
        import trafilatura
        # Use urllib request to avoid python signal errors inside threads
        with open_checked(
            url,
            headers={"User-Agent": "Mozilla/5.0"},
            timeout=15,
        ) as r:
            downloaded = r.read(5 * 1024 * 1024 + 1)
        if not downloaded or len(downloaded) > 5 * 1024 * 1024:
            return {"error": "fetch_failed"}
        from trafilatura.settings import use_config
        newconfig = use_config()
        newconfig.set("DEFAULT", "EXTRACTION_TIMEOUT", "0")
        result = trafilatura.extract(
            downloaded,
            output_format="json",
            with_metadata=True,
            include_comments=False,
            include_images=False,
            favor_recall=True,
            config=newconfig,
        )
        if not result:
            return {"error": "extract_failed"}
        data = json.loads(result)
        text = data.get("text", "") or ""
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
        image = data.get("image")
        # запасной поиск og:image. downloaded — bytes (urllib), декодируем для
        # regex; ошибка здесь не должна ронять всю статью.
        if not image:
            try:
                html_str = downloaded.decode("utf-8", "ignore") \
                    if isinstance(downloaded, (bytes, bytearray)) else downloaded
                m = re.search(
                    r'<meta[^>]+property=["\']og:image["\'][^>]+content=["\']([^"\']+)["\']',
                    html_str, re.IGNORECASE)
                if m:
                    image = m.group(1)
            except Exception:
                pass
        result = {
            "title": data.get("title", "") or "",
            "image": image,
            "source": data.get("sitename") or data.get("hostname", "") or "",
            "date": data.get("date", "") or "",
            "paragraphs": paragraphs,
        }
        if paragraphs:
            redis_cache.set_json(f"article:{cache_key}", result, 6 * 3600)
        return result
    except Exception as ex:
        logging.error(f"Article extraction failed ({url}): {ex}")
        return {"error": "fetch_failed"}


# ---------------------------------------------------------------------------
# Аудирование (бета): TTS-озвучка предложений + курируемые уроки с субтитрами.
# ---------------------------------------------------------------------------
import json as _json
from fastapi.responses import Response, StreamingResponse

_TTS_CACHE: Dict[str, bytes] = {}
_TTS_CACHE_MAX = 600  # ~ десятки минут речи; примитивный FIFO достаточен


@app.get("/audio/tts")
def audio_tts(
    text: str,
    lang: str = "sr",
    voice: str = "sophie",
    _: None = Depends(_tts_limit),
):
    """Озвучивает короткий текст нейросетевым голосом, отдаёт mp3.

    Клиент шлёт по одному предложению — так у него точные тайминги реплик
    (каждая реплика = свой файл) для караоке-подсветки.
    """
    if lang not in ("sr", "ru", "en"):
        raise HTTPException(status_code=400, detail="Неподдерживаемый язык.")
    if voice not in ("sophie", "nicholas"):
        raise HTTPException(status_code=400, detail="Неподдерживаемый диктор.")
    text = (text or "").strip()
    if not text or len(text) > 400:
        return Response(status_code=400)
    serbian_name = "SophieNeural" if voice == "sophie" else "NicholasNeural"
    voices = {
        # edge-tts 7.x отклоняет трёхчастный locale sr-Latn-RS ещё до
        # запроса, хотя Azure публикует такой alias. sr-RS принимает обе
        # сербские письменности и не сваливается на роботизированный gTTS.
        "sr": f"sr-RS-{serbian_name}",
        "ru": "ru-RU-SvetlanaNeural",
        "en": "en-US-JennyNeural",
    }
    voice = voices[lang]
    key = hashlib.sha1(f"neural-v2:{voice}:{text}".encode("utf-8")).hexdigest()
    data = _TTS_CACHE.get(key)
    if data is None:
        data = redis_cache.get_bytes(f"tts:{key}")
    if data is None:
        try:
            import edge_tts
            chunks = []
            for message in edge_tts.Communicate(text, voice, rate="-4%").stream_sync():
                if message["type"] == "audio":
                    chunks.append(message["data"])
            data = b"".join(chunks)
            if not data:
                raise RuntimeError("neural TTS returned no audio")
        except Exception as neural_error:
            logging.warning(f"Neural TTS failed, falling back to gTTS: {neural_error}")
            try:
                from io import BytesIO
                from gtts import gTTS
                buf = BytesIO()
                gTTS(text=text, lang=lang).write_to_fp(buf)
                data = buf.getvalue()
            except Exception as fallback_error:
                logging.error(f"TTS failed: {fallback_error}")
                return Response(status_code=502)
        if len(_TTS_CACHE) >= _TTS_CACHE_MAX:
            _TTS_CACHE.pop(next(iter(_TTS_CACHE)))
        _TTS_CACHE[key] = data
        redis_cache.set_bytes(f"tts:{key}", data, 7 * 24 * 3600)
    elif key not in _TTS_CACHE:
        if len(_TTS_CACHE) >= _TTS_CACHE_MAX:
            _TTS_CACHE.pop(next(iter(_TTS_CACHE)))
        _TTS_CACHE[key] = data
    return Response(
        content=data,
        media_type="audio/mpeg",
        headers={"Cache-Control": "public, max-age=86400"},
    )


# Подкасты на сербском с открытым RSS (прямые mp3 в enclosure).
#  - Može Kafa: anchor-фид, проверен (46 эпизодов, mp3 с cloudfront);
#  - Learn Serbian Podcast (serbianlanguagelessons.com): buzzsprout-фид из
#    официального каталога iTunes.
# Текст реплик берём из описания эпизода (у этих подкастов туда кладут
# транскрипт/конспект на сербском) и растягиваем пропорционально длительности.
PODCAST_FEEDS = [
    ("learn-serbian", "Learn Serbian Podcast",
     "https://rss.buzzsprout.com/1246415.rss"),
    ("moze-kafa", "Može Kafa Podcast",
     "https://anchor.fm/s/aef64434/podcast/rss"),
]
_PODCAST_CACHE: Dict[str, object] = {"ts": 0.0, "items": []}
_PODCAST_TTL = 1800  # 30 минут
_AUDIO_PROXY_HOSTS = {
    "www.buzzsprout.com",
    "buzzsprout.com",
    "anchor.fm",
    "audio.buzzsprout.com",
}


def _duration_seconds(raw) -> float:
    """itunes:duration бывает '1234' или '00:23:55'."""
    try:
        s = str(raw or "").strip()
        if not s:
            return 0.0
        if ":" in s:
            parts = [int(p) for p in s.split(":")]
            sec = 0
            for p in parts:
                sec = sec * 60 + p
            return float(sec)
        return float(s)
    except Exception:
        return 0.0


def _split_sentences(text: str):
    import re as _re
    out = []
    # Границы: конец предложения или перенос строки (из <br>/<p> разметки).
    for s in _re.split(r"(?<=[.!?…])\s+|\n+", text):
        s = s.strip()
        if len(s) < 2:
            continue
        # очень длинные куски режем по запятой, чтобы реплики были обозримыми
        while len(s) > 260:
            cut = s.rfind(",", 0, 240)
            if cut < 80:
                cut = 240
            out.append(s[: cut + 1].strip())
            s = s[cut + 1:].strip()
        if s:
            out.append(s)
    return out


def _proportional_cues(sents, dur: float):
    """Реплики с таймингами, растянутыми пропорционально длине текста."""
    total_chars = sum(len(s) for s in sents) or 1
    cues, t0 = [], 0.0
    for s in sents:
        if dur > 0:
            t1 = t0 + dur * len(s) / total_chars
            cues.append({"start": round(t0, 2), "end": round(t1, 2), "text": s})
            t0 = t1
        else:
            cues.append({"text": s})
    return cues


def _extract_transcript_text(html_str: str) -> str:
    """Достаёт читаемый текст страницы транскрипта без trafilatura.

    trafilatura внутри FastAPI threadpool может падать с `signal only works in
    main thread`, а Wix-страницы serbianlanguagelessons.com уже содержат текст
    транскрипта в HTML. Поэтому здесь достаточно аккуратной stdlib-очистки.
    """
    text = re.sub(r"<script[^>]*>.*?</script>", " ", html_str,
                  flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"<style[^>]*>.*?</style>", " ", text,
                  flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"<(noscript|svg)[^>]*>.*?</\1>", " ", text,
                  flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(
        r"<\s*(br|/p|/div|/li|/h[1-6]|/section|/article)[^>]*>",
        "\n",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"<[^>]+>", " ", text)
    text = _html.unescape(text)
    text = re.sub(
        r"\\u([0-9a-fA-F]{4})",
        lambda m: chr(int(m.group(1), 16)),
        text,
    )

    boilerplate = re.compile(
        r"(top of page|log in|free materials|plans & pricing|"
        r"blog & transcripts|about the tutor|all posts|search|"
        r"recent posts|comments|write a comment|bottom of page|"
        r"privacy policy|terms|subscribe)",
        re.IGNORECASE,
    )
    lines = []
    for line in text.splitlines():
        line = re.sub(r"\s+", " ", line).strip()
        if len(line) < 2 or boilerplate.search(line):
            continue
        lines.append(line)

    # Отрезаем шапку страницы до вероятного начала транскрипта, но не
    # требуем строгого маркера: у разных выпусков начало оформлено по-разному.
    start = 0
    for i, line in enumerate(lines[:80]):
        if re.search(r"\b(Zdravo|Dobro|Dobar|U redu|Hajde|Welcome)\b",
                     line, re.IGNORECASE):
            start = i
            break
    return "\n".join(lines[start:])[:60000]


@app.get("/audio/proxy")
def audio_proxy(url: str, request: Request):
    """Прокси для web-аудио из RSS.

    У части podcast/CDN-хостов финальный mp3/m4a не отдаёт CORS-заголовки или
    режет запросы с Origin, из-за чего HTMLAudioElement получает HTML-ошибку
    вместо аудио. Проксируем только известные podcast-хосты и пробрасываем Range.
    """
    parsed = urllib.parse.urlparse(url)
    host = (parsed.hostname or "").lower()
    if parsed.scheme not in {"https", "http"} or host not in _AUDIO_PROXY_HOSTS:
        return Response(status_code=403)

    headers = {
        "User-Agent": "Mozilla/5.0",
        "Accept": "*/*",
    }
    range_header = request.headers.get("range")
    if range_header:
        headers["Range"] = range_header
    if host.endswith("buzzsprout.com"):
        headers["Referer"] = "https://www.buzzsprout.com/"
    elif host == "anchor.fm":
        headers["Referer"] = "https://anchor.fm/"

    try:
        req = urllib.request.Request(url, headers=headers)
        upstream = urllib.request.urlopen(req, timeout=30)
    except Exception as ex:
        logging.error(f"Audio proxy failed ({url}): {ex}")
        return Response(status_code=502)

    status = getattr(upstream, "status", 200)
    content_type = upstream.headers.get("Content-Type") or "application/octet-stream"
    response_headers = {
        "Content-Type": content_type,
        "Accept-Ranges": upstream.headers.get("Accept-Ranges", "bytes"),
        "Cache-Control": "public, max-age=3600",
        "Access-Control-Allow-Origin": "*",
        "Access-Control-Expose-Headers": "Accept-Ranges, Content-Length, Content-Range",
    }
    for name in ("Content-Length", "Content-Range"):
        value = upstream.headers.get(name)
        if value:
            response_headers[name] = value

    def body():
        try:
            while True:
                chunk = upstream.read(64 * 1024)
                if not chunk:
                    break
                yield chunk
        finally:
            upstream.close()

    return StreamingResponse(body(), status_code=status, headers=response_headers)


TRANSCRIPTS_BASE = "https://citavuk.ru/transcripts"
_TRANSCRIPT_INDEX: Dict[str, object] = {"ts": 0.0, "map": {}}
_TRANSCRIPT_INDEX_TTL = 3600


def _transcript_index() -> dict:
    """Карта «ссылка на аудио → имя файла с расшифровкой».

    Файлы делает web/scripts/transcribe-podcasts.py и выкладывает вместе с
    сайтом; здесь только индекс, чтобы не дёргать сеть на каждый эпизод.
    """
    import time as _t
    now = _t.time()
    if now - float(_TRANSCRIPT_INDEX["ts"]) < _TRANSCRIPT_INDEX_TTL:
        return _TRANSCRIPT_INDEX["map"]  # type: ignore[return-value]
    try:
        req = urllib.request.Request(
            f"{TRANSCRIPTS_BASE}/index.json",
            headers={"User-Agent": "citavuk"})
        data = json.loads(urllib.request.urlopen(req, timeout=15).read())
        if isinstance(data, dict):
            _TRANSCRIPT_INDEX["map"] = data
            _TRANSCRIPT_INDEX["ts"] = now
    except Exception as ex:
        logging.warning(f"Transcript index unavailable: {ex}")
        _TRANSCRIPT_INDEX["ts"] = now
    return _TRANSCRIPT_INDEX["map"]  # type: ignore[return-value]


def _own_transcript_url(audio_url: str) -> Optional[str]:
    name = _transcript_index().get((audio_url or "").split("?")[0])
    return f"{TRANSCRIPTS_BASE}/{name}" if name else None


def _check_transcript_url(url: str) -> str:
    """Разрешает только два источника и сохраняет правило на редиректах."""
    candidate = (url or "").strip()
    parsed = urllib.parse.urlparse(candidate)
    host = (parsed.hostname or "").lower()
    if parsed.scheme not in ("http", "https"):
        raise UrlRejected("недопустимая схема транскрипта")
    if host == "citavuk.ru" and parsed.path.startswith("/transcripts/"):
        return candidate
    if host == "www.serbianlanguagelessons.com":
        return candidate
    raise UrlRejected("источник транскрипта вне списка")


@app.get("/audio/transcript")
def audio_transcript(url: str, duration: float = 0.0):
    """Полный транскрипт эпизода.

    Домен ограничен — это не открытый прокси.
    """
    try:
        _check_transcript_url(url)
    except UrlRejected:
        return {"cues": []}
    parsed = urllib.parse.urlparse(url)
    host = (parsed.hostname or "").lower()
    if host == "citavuk.ru":
        try:
            with open_checked(
                url,
                headers={"User-Agent": "citavuk"},
                timeout=25,
                validator=_check_transcript_url,
            ) as response:
                data = json.loads(response.read())
            cues = data.get("cues")
            return {"cues": cues if isinstance(cues, list) else []}
        except Exception as ex:
            logging.error(f"Own transcript failed ({url}): {ex}")
            return {"cues": []}
    cache_key = hashlib.sha256(
        f"{url.strip()}:{round(duration, 2)}".encode("utf-8")
    ).hexdigest()
    cached = redis_cache.get_json(f"transcript:{cache_key}")
    if isinstance(cached, dict) and isinstance(cached.get("cues"), list):
        return cached
    try:
        with open_checked(
            url,
            headers={"User-Agent": "Mozilla/5.0"},
            timeout=25,
            validator=_check_transcript_url,
        ) as response:
            raw = response.read()
        html_str = raw.decode("utf-8", "ignore") if isinstance(
            raw, (bytes, bytearray)) else raw
        text = _extract_transcript_text(html_str)
        sents = [
            s for s in _split_sentences(text)
            if "http" not in s and not s.startswith("#")
        ]
        result = {"cues": _proportional_cues(sents, duration)}
        if result["cues"]:
            redis_cache.set_json(
                f"transcript:{cache_key}", result, 24 * 3600
            )
        return result
    except Exception as ex:
        logging.error(f"Transcript extraction failed ({url}): {ex}")
        return {"cues": []}


def _podcast_lessons(limit_per_feed: int = 12):
    import time as _t
    now = _t.time()
    if now - _PODCAST_CACHE["ts"] < _PODCAST_TTL and _PODCAST_CACHE["items"]:
        return _PODCAST_CACHE["items"]
    shared = redis_cache.get_json("podcasts")
    if isinstance(shared, list) and shared:
        _PODCAST_CACHE["ts"] = now
        _PODCAST_CACHE["items"] = shared
        return shared
    import feedparser
    items = []
    for fid, fname, url in PODCAST_FEEDS:
        try:
            feed = feedparser.parse(url)
            for e in feed.entries[:limit_per_feed]:
                audio = None
                for l in e.get("links", []):
                    if "audio" in (l.get("type") or ""):
                        audio = l.get("href")
                        break
                if not audio and e.get("enclosures"):
                    audio = e.enclosures[0].get("href")
                if not audio:
                    continue
                # Полная очистка html (без усечения — _clean_summary режет
                # до 300 симв., а тут транскрипты). <br>/<p> → границы фраз.
                raw = e.get("summary", "") or ""
                raw = re.sub(r"<\s*(br|/p|/div|/li)[^>]*>", "\n", raw)
                text = re.sub(r"<[^>]+>", " ", raw)
                text = _html.unescape(text)
                text = re.sub(r"[ \t]+", " ", text)
                text = "\n".join(
                    ln.strip() for ln in text.splitlines() if ln.strip())
                # buzzsprout добавляет мусорные строки («Send a text»,
                # ссылка на транскрипт) — в реплики они не нужны.
                sents = [
                    s for s in _split_sentences(text)
                    if "http" not in s and "Send a text" not in s
                ]
                if not sents:
                    continue
                dur = _duration_seconds(e.get("itunes_duration"))
                cues = _proportional_cues(sents, dur)
                mins = f" · {int(dur // 60)} мин" if dur > 0 else ""
                item = {
                    "id": f"{fid}-{e.get('id', e.get('title', ''))}"[:120],
                    "title": e.get("title", "") or "Без названия",
                    "subtitle": f"{fname}{mins}",
                    "audio_url": audio,
                    "cues": cues,
                }
                # Расшифровка самого аудио, если она уже сделана: текст со
                # страницы подкаста расходился с речью, а тайминги там были
                # выдуманы пропорционально длительности.
                own = _own_transcript_url(audio)
                if own:
                    item["transcript_url"] = own
                    item["duration"] = dur
                else:
                    # Запасной вариант — транскрипт со страницы автора.
                    m = re.search(
                        r"https://www\.serbianlanguagelessons\.com/post/[\w%().~-]+",
                        _html.unescape(e.get("summary", "") or ""))
                    if m:
                        item["transcript_url"] = m.group(0)
                        item["duration"] = dur
                items.append(item)
        except Exception as ex:
            logging.error(f"Podcast feed failed ({url}): {ex}")
    if items:
        _PODCAST_CACHE["ts"] = now
        _PODCAST_CACHE["items"] = items
        redis_cache.set_json("podcasts", items, _PODCAST_TTL)
    return items


@app.get("/audio/lessons")
def audio_lessons():
    """Аудиоуроки: курируемый audio_lessons.json + эпизоды подкастов из RSS.

    audio_lessons.json лежит рядом с main.py и редактируется прямо на Space
    без пересборки приложения. Формат:
    {"items": [{"id": "...", "title": "...", "subtitle": "...",
                "audio_url": "https://...mp3",
                "cues": [{"start": 0.0, "end": 4.2, "text": "..."}, ...]}]}
    """
    curated = []
    for path in ("audio_lessons.json", "../audio_lessons.json"):
        if os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = _json.load(f)
                curated = data.get("items", [])
            except Exception as e:
                logging.error(f"audio_lessons.json parse failed: {e}")
            break
    return {"items": curated + _podcast_lessons()}


if __name__ == "__main__":
    import uvicorn
    # Run backend
    uvicorn.run(app, host="0.0.0.0", port=8000)
